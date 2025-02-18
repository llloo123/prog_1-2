import tkinter as tk
from tkinter import messagebox
from tkinter.filedialog import asksaveasfilename
from vehicles.car import Car
from vehicles.truck import Truck
from vehicles.bus import Bus
from openpyxl import Workbook
from docx import Document

def calculate():
    vehicle_type = vehicle_var.get()
    try:
        distance = float(distance_entry.get())
        load = float(load_entry.get())
        fuel_price = float(fuel_price_entry.get())

        if vehicle_type == "Car":
            vehicle = Car(distance, load, fuel_price)
        elif vehicle_type == "Truck":
            vehicle = Truck(distance, load, fuel_price)
        elif vehicle_type == "Bus":
            vehicle = Bus(distance, load, fuel_price)
        else:
            raise ValueError("Invalid vehicle type")

        fuel_consumption = vehicle.calculate_fuel_consumption()
        cost = vehicle.calculate_cost()
        time = vehicle.calculate_time()

        result_text.set(f"Fuel Consumption: {fuel_consumption:.2f} liters\nCost: ${cost:.2f}\nTime: {time:.2f} hours")
    except Exception as e:
        messagebox.showerror("Error", str(e))

def save_results():
    vehicle_type = vehicle_var.get()
    try:
        distance = float(distance_entry.get())
        load = float(load_entry.get())
        fuel_price = float(fuel_price_entry.get())

        if vehicle_type == "Car":
            vehicle = Car(distance, load, fuel_price)
        elif vehicle_type == "Truck":
            vehicle = Truck(distance, load, fuel_price)
        elif vehicle_type == "Bus":
            vehicle = Bus(distance, load, fuel_price)
        else:
            raise ValueError("Invalid vehicle type")

        fuel_consumption = vehicle.calculate_fuel_consumption()
        cost = vehicle.calculate_cost()
        time = vehicle.calculate_time()

        file_path = asksaveasfilename(defaultextension=".doc", filetypes=[("Word Documents", "*.doc"), ("Excel Files", "*.xls")])
        if not file_path:
            return

        if file_path.endswith(".doc"):
            doc = Document()
            doc.add_heading(f"{vehicle_type} Trip Details", level=1)
            doc.add_paragraph(f"Distance: {distance} km")
            doc.add_paragraph(f"Load: {load}%")
            doc.add_paragraph(f"Fuel Price: ${fuel_price}")
            doc.add_paragraph(f"Fuel Consumption: {fuel_consumption:.2f} liters")
            doc.add_paragraph(f"Cost: ${cost:.2f}")
            doc.add_paragraph(f"Time: {time:.2f} hours")
            doc.save(file_path)
        elif file_path.endswith(".xls"):
            wb = Workbook()
            ws = wb.active
            ws.title = f"{vehicle_type} Trip Details"
            ws.append(["Property", "Value"])
            ws.append(["Distance", f"{distance} km"])
            ws.append(["Load", f"{load}%"])
            ws.append(["Fuel Price", f"${fuel_price}"])
            ws.append(["Fuel Consumption", f"{fuel_consumption:.2f} liters"])
            ws.append(["Cost", f"${cost:.2f}"])
            ws.append(["Time", f"{time:.2f} hours"])
            wb.save(file_path)

        messagebox.showinfo("Success", "Results saved successfully!")
    except Exception as e:
        messagebox.showerror("Error", str(e))

root = tk.Tk()
root.title("Vehicle Calculator")

vehicle_var = tk.StringVar(value="Car")
tk.Label(root, text="Select Vehicle Type:").grid(row=0, column=0, sticky="w")
tk.OptionMenu(root, vehicle_var, "Car", "Truck", "Bus").grid(row=0, column=1, sticky="w")

distance_label = tk.Label(root, text="Distance (km):")
distance_entry = tk.Entry(root)
load_label = tk.Label(root, text="Load (%):")
load_entry = tk.Entry(root)
fuel_price_label = tk.Label(root, text="Fuel Price ($/liter):")
fuel_price_entry = tk.Entry(root)

result_text = tk.StringVar()
result_label = tk.Label(root, textvariable=result_text, justify="left")

calculate_button = tk.Button(root, text="Calculate", command=calculate)
save_button = tk.Button(root, text="Save Results", command=save_results)

tk.Label(root, text="Parameters:").grid(row=1, column=0, sticky="w")
distance_label.grid(row=2, column=0, sticky="w")
distance_entry.grid(row=2, column=1, sticky="w")
load_label.grid(row=3, column=0, sticky="w")
load_entry.grid(row=3, column=1, sticky="w")
fuel_price_label.grid(row=4, column=0, sticky="w")
fuel_price_entry.grid(row=4, column=1, sticky="w")
calculate_button.grid(row=5, column=0, columnspan=2, pady=10)
save_button.grid(row=6, column=0, columnspan=2, pady=10)
result_label.grid(row=7, column=0, columnspan=2, pady=10)

root.mainloop()